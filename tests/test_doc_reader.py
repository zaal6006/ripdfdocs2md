import subprocess
from pathlib import Path

import pytest
from docx import Document

from ripdfdocs2md import doc_reader
from ripdfdocs2md.pipeline import convert_file

try:
    _SOFFICE = doc_reader.find_soffice()
except doc_reader.SofficeNotFoundError:
    _SOFFICE = None

pytestmark = pytest.mark.skipif(
    _SOFFICE is None,
    reason="LibreOffice (soffice) not found - set RIPDFDOCS2MD_SOFFICE or install it to run .doc tests",
)


def _make_doc(tmp_path: Path) -> Path:
    """Build a real legacy .doc file for testing by round-tripping a
    synthetic .docx through LibreOffice itself — python-docx can't write
    the old binary format directly, so this is the only way to get a
    genuine .doc fixture without committing a real document to the repo.
    """
    docx_path = tmp_path / "source.docx"
    document = Document()
    document.add_heading("Test Heading", level=1)
    document.add_paragraph("Hello from a test DOC.")
    document.save(docx_path)

    profile_dir = tmp_path / "lo_profile"
    subprocess.run(
        [
            _SOFFICE,
            "--headless",
            "--norestore",
            f"-env:UserInstallation={profile_dir.as_uri()}",
            "--convert-to",
            "doc:MS Word 97",
            "--outdir",
            str(tmp_path),
            str(docx_path),
        ],
        capture_output=True,
        text=True,
        timeout=120,
        check=True,
    )
    return tmp_path / "source.doc"


def test_doc_reader_preserves_structure(tmp_path):
    doc_path = _make_doc(tmp_path)

    markdown = doc_reader.convert(doc_path)

    assert "Test Heading" in markdown
    assert "Hello from a test DOC" in markdown


def test_convert_file_routes_doc_through_pipeline(tmp_path):
    doc_path = _make_doc(tmp_path)

    markdown = convert_file(doc_path)

    assert "Hello from a test DOC" in markdown


def test_find_soffice_rejects_bad_override(monkeypatch):
    monkeypatch.setenv("RIPDFDOCS2MD_SOFFICE", str(Path("nonexistent") / "soffice.exe"))

    with pytest.raises(doc_reader.SofficeNotFoundError):
        doc_reader.find_soffice()
